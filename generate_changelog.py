"""
26/04/19 수정 사항 정리 Word 문서 생성 스크립트.
python generate_changelog.py 로 실행하면 같은 폴더에 changelog_26_04_19.docx 생성.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(11)
    return p


def add_code(doc, code):
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = "Menlo"
    r.font.size = Pt(9)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def main():
    doc = Document()

    # 제목
    title = doc.add_heading("실시간 수어 번역 백엔드 - 26/04/19 수정 정리", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "작성일: 2026-04-19 / 대상: 실시간 수어 번역 서비스 백엔드 "
        "(Real-time Sign Language Translation Service Backend)",
        italic=True,
    )
    doc.add_paragraph()

    # ================================================================
    # 1. 환경 설정
    # ================================================================
    add_heading(doc, "1. 환경 설정 수정", 1)

    add_heading(doc, "1.1 .env 파일 누락 변수 추가", 2)
    add_para(
        doc,
        "Settings (pydantic-settings) 초기화 시 JWT/Kakao/Redis 관련 7개 필수 변수가 "
        "누락되어 서버 기동이 실패했음. .env 파일에 다음 항목 추가:",
    )
    add_code(
        doc,
        """# JWT 설정
JWT_SECRET_KEY=asdfas234edsf234f
JWT_ALGORITHM=HS256
ACCESS_TOKEN_EXPIRE_MINUTES=30

# Kakao API 설정
KAKAO_REST_API_KEY=...
KAKAO_REDIRECT_URI=http://localhost:8080/api/v1/auth/kakao/auth
KAKAO_CLIENT_SECRET=...

# Redis 설정
REDIS_HOST=127.0.0.1
REDIS_PORT=6379""",
    )

    add_heading(doc, "1.2 PyJWT 설치", 2)
    add_para(
        doc,
        "security.py 가 `import jwt` (PyJWT) 를 사용하는데 requirements.txt 에 "
        "python-jose 만 있어서 누락. PyJWT==2.9.0 을 requirements.txt 에 추가.",
    )
    add_para(
        doc,
        "주의: pip install jwt 는 다른 무명 패키지이므로 설치 시 pyjwt 또는 PyJWT 로 정확히 지정 필요.",
        italic=True,
    )

    add_heading(doc, "1.3 requirements.txt ML 패키지 버전 고정", 2)
    add_para(
        doc,
        "Python 3.13 의 mediapipe 최신 버전에는 mp.solutions 가 제거되어 지문자 인식이 불가. "
        "팀 공용으로 Python 3.11 환경(py311_env) 기준으로 버전 pin:",
    )
    add_code(
        doc,
        """torch==2.9.1
mediapipe==0.10.21
opencv-python==4.11.0.86
Pillow==12.0.0
numpy==1.26.4""",
    )

    # ================================================================
    # 2. 로그인 플로우 연결
    # ================================================================
    add_heading(doc, "2. 로그인 플로우 연결 (프론트엔드 연동)", 1)

    add_heading(doc, "2.1 카카오 로그인 버튼 핸들러 (신규)", 2)
    add_para(doc, "파일: frontend/js/login.js (NEW)")
    add_para(
        doc,
        "로그인 버튼 클릭 시 /api/v1/auth/login/kakao 로 이동 — "
        "서버가 카카오 OAuth URL 로 리다이렉트.",
    )
    add_code(
        doc,
        """document.addEventListener('DOMContentLoaded', () => {
    const kakaoBtn = document.querySelector('.kakao-login-btn');
    if (kakaoBtn) {
        kakaoBtn.addEventListener('click', () => {
            window.location.href = '/api/v1/auth/login/kakao';
        });
    }
});""",
    )

    add_heading(doc, "2.2 로그인 콜백에 access_token 쿠키 추가", 2)
    add_para(doc, "파일: main/api/user/user_routes.py")
    add_para(
        doc,
        "refresh_token 은 httpOnly 쿠키로 저장돼 있었지만 access_token 은 응답 헤더에만 있어서 "
        "JS 가 접근 불가. 프론트엔드가 Authorization 헤더를 직접 만들기 위해 "
        "access_token 을 non-httpOnly 쿠키로도 저장:",
    )
    add_code(
        doc,
        """response.set_cookie(
    key="access_token",
    value=my_access_token,
    httponly=False,       # JS 가 읽을 수 있어야 함
    secure=False,
    samesite="lax",
    max_age=settings.ACCESS_TOKEN_EXPIRE_MINUTES * 60,
)""",
    )

    add_heading(doc, "2.3 가입하기 폼 핸들러 (신규)", 2)
    add_para(doc, "파일: frontend/js/register.js (NEW), frontend/html/register.html 에 include 추가")
    add_para(
        doc,
        "기존에는 <form> submit 기본 동작만 있었고 API 호출 JS 가 없었음. "
        "쿠키에서 access_token 읽어 Authorization 헤더로 POST /api/v1/users/info 호출. "
        "성공 시 home.html 로 이동.",
    )

    # ================================================================
    # 3. 학습 라우팅 / UI 정리
    # ================================================================
    add_heading(doc, "3. 학습(Learning) 라우팅 및 UI 정리", 1)

    add_heading(doc, "3.1 API_BASE 경로 수정", 2)
    add_para(
        doc,
        "frontend/js/learning.js, frontend/js/sign_learn.js 에서 /api/learning 을 "
        "/api/v1/learning 으로 변경. 백엔드 라우터 등록 prefix 와 일치시킴.",
    )

    add_heading(doc, "3.2 DB 시드 초기화 버튼 제거", 2)
    add_para(
        doc,
        "learning.html 에서 'DB 시드 초기화' 버튼 UI 삭제, learning.js 에서 seed() 함수와 "
        "클릭 핸들러 제거. 백엔드 /seed/fingerspell 엔드포인트는 유지 (관리자 용도).",
    )

    add_heading(doc, "3.3 /logo.png 정적 파일 라우트 추가", 2)
    add_para(doc, "파일: main/main.py")
    add_para(
        doc,
        "HTML 이 <img src='../logo.png'> 로 참조하는데 StaticFiles mount 에 logo 경로가 "
        "없어서 404 발생. 단일 라우트 추가:",
    )
    add_code(
        doc,
        """@app.get("/logo.png")
def serve_logo():
    return FileResponse(frontend_dir / "logo.png")""",
    )

    # ================================================================
    # 4. 학습 결과 DB 저장
    # ================================================================
    add_heading(doc, "4. 학습 결과 DB 저장 (user_lesson_progress 연결)", 1)

    add_heading(doc, "4.1 Repository 신규 작성", 2)
    add_para(
        doc,
        "파일: main/domain/UserLessonProgress/repository/user_lesson_progress_repository.py (NEW)",
    )
    add_para(
        doc,
        "find_by_user_and_lesson, save 두 메서드를 가진 Sync Repository. "
        "같은 (user_id, lesson_id) 조합으로 조회 후 insert 혹은 update 분기.",
    )

    add_heading(doc, "4.2 SaveResultUseCase upsert 로 변경", 2)
    add_para(doc, "파일: main/domain/learning/usecase/lesson_usecase.py")
    add_para(
        doc,
        "기존에는 판정만 하고 DB 저장 없이 응답 반환. user_id 를 인자로 받고 "
        "진행도 테이블에 upsert. status: 'passed' (80점 이상) 혹은 'failed'. "
        "한 번 passed 로 저장된 이후에는 status 덮어쓰지 않음(진행 유지).",
    )

    add_heading(doc, "4.3 /results 엔드포인트에 JWT 필수화", 2)
    add_para(doc, "파일: main/api/learning/learning_routes.py")
    add_para(
        doc,
        "Depends(get_current_user_id) 추가. 프론트엔드에서 Authorization 헤더 없이 "
        "요청하면 401 반환.",
    )

    add_heading(doc, "4.4 프론트엔드 Authorization 헤더 전송", 2)
    add_para(doc, "파일: frontend/js/sign_learn.js, frontend/js/word_learn.js")
    add_para(
        doc,
        "getCookie('access_token') 으로 토큰 획득 후 Bearer 헤더에 담아 POST /results. "
        "401 응답 시 login.html 로 이동.",
    )

    # ================================================================
    # 5. 단어 시드 + 카테고리
    # ================================================================
    add_heading(doc, "5. 단어(Word) 시드 및 카테고리 분류", 1)

    add_heading(doc, "5.1 단어 리스트 추가 (약 398개)", 2)
    add_para(doc, "파일: main/domain/learning/service/lesson_service.py")
    add_para(
        doc,
        "Excel 원본에서 받은 단어를 10개 카테고리로 분류해 WORDS_BY_CATEGORY 딕셔너리로 정의. "
        "동일 단어가 여러 카테고리에 등장하면 먼저 나온 카테고리 우선 규칙으로 "
        "WORD_CATEGORY_MAP (단어→카테고리) 평탄화.",
    )
    add_para(doc, "카테고리 10종:", bold=True)
    for cat in [
        "greeting (인사·기본 표현)",
        "family (사람·가족·관계)",
        "emotion (감정·기분)",
        "body (신체·건강)",
        "job (직업)",
        "nature (자연·날씨·계절)",
        "thing (사물·음식·기기)",
        "action (동작·상태·일상)",
        "place (장소·시간)",
        "study (학업·업무·숫자·기타)",
    ]:
        add_bullet(doc, cat)

    add_heading(doc, "5.2 SeedWordsUseCase: insert + update 지원", 2)
    add_para(
        doc,
        "기존 row 의 subcategory 가 NULL 이거나 다르면 UPDATE, 없으면 INSERT. "
        "응답 메시지에 'X개 추가, Y개 카테고리 업데이트, Z개 스킵' 표기.",
    )

    add_heading(doc, "5.3 POST /api/v1/learning/seed/word 엔드포인트 추가", 2)
    add_para(doc, "파일: main/api/learning/learning_routes.py")

    # ================================================================
    # 6. 지문자 확장
    # ================================================================
    add_heading(doc, "6. 지문자 확장 및 모델 라우팅 분기", 1)

    add_heading(doc, "6.1 지문자 시드에 확장 자모 추가", 2)
    add_para(doc, "파일: main/domain/learning/service/lesson_service.py")
    add_para(
        doc,
        "KOREAN_CONSONANTS 에 쌍자음 5개 (ㄲ ㄸ ㅃ ㅆ ㅉ) 추가. "
        "KOREAN_VOWELS 에 이중모음 4개 (ㅘ ㅙ ㅝ ㅞ) + fingerspell 모델이 지원하는 "
        "추가 모음 3개 (ㅚ ㅟ ㅢ) 추가. 총 37개 (기존 28 + 9).",
    )

    add_heading(doc, "6.2 모델별 라우팅 분기", 2)
    add_para(doc, "파일: frontend/js/learning.js")
    add_para(
        doc,
        "model_fingerspell.pt 의 라벨 31개는 sign_learn.html (fingerspell 모델 WS) 로 이동. "
        "model_fingerspell.pt 에 없고 model_word.pt 에만 있는 9개 "
        "(ㄲ ㄸ ㅃ ㅆ ㅉ ㅘ ㅙ ㅝ ㅞ) 는 word_learn.html 로 이동하여 "
        "word 모델로 인식.",
    )
    add_code(
        doc,
        """const WORD_MODEL_FS_CHARS = new Set([
  "ㄲ", "ㄸ", "ㅃ", "ㅆ", "ㅉ", "ㅘ", "ㅙ", "ㅝ", "ㅞ"
]);
const page = WORD_MODEL_FS_CHARS.has(item.title)
    ? "word_learn.html"
    : "sign_learn.html";""",
    )

    # ================================================================
    # 7. 단어 학습 UI (Phase 1-2)
    # ================================================================
    add_heading(doc, "7. 단어 학습 UI 구축 (Phase 1-2)", 1)

    add_heading(doc, "7.1 learning.html 에 단어 학습 섹션 추가", 2)
    add_para(
        doc,
        "지문자 학습 섹션 아래에 '단어 학습' 제목, 카테고리 탭 DOM (#wordTabs), "
        "단어 카드 그리드 (#wordList) 배치.",
    )

    add_heading(doc, "7.2 learning.css 단어 섹션 스타일", 2)
    add_para(
        doc,
        ".word-tabs (탭 래퍼), .word-tab (탭 버튼, active 상태 포함), "
        ".lesson-card.word-card (단어용 카드: 정사각형 아닌 직사각형, 글자 중심).",
    )

    add_heading(doc, "7.3 learning.js 카테고리 렌더링", 2)
    add_para(
        doc,
        "WORD_CATEGORIES 배열로 탭 생성 (10개). 최초 'greeting' 활성. "
        "탭 클릭 시 activeCategory 변경 후 해당 subcategory 단어만 카드로 렌더. "
        "GET /lessons?category=word 한 번 호출로 전체 받아 클라이언트 측 필터링.",
    )

    add_heading(doc, "7.4 word_learn.html / word_learn.js (신규)", 2)
    add_para(
        doc,
        "지문자 학습(sign_learn)의 4단계 Stepper 구조를 복제하되 단어용으로 "
        "수정. Step 1 단어 표시, Step 2 카메라 테스트, Step 3 실제 인식, Step 4 완료.",
    )

    # ================================================================
    # 8. 단어 인식 (Phase 3)
    # ================================================================
    add_heading(doc, "8. 단어 실시간 인식 (Phase 3, model_word.pt 연결)", 1)

    add_heading(doc, "8.1 WordRecognitionService 신규", 2)
    add_para(
        doc,
        "파일: main/domain/learning/service/word_recognition_service.py (NEW)",
    )
    add_para(doc, "recon_word.py 의 로직을 웹 환경으로 포팅:")
    add_bullet(doc, "FrameMLP 모델 (114차원 손 특징 → 915 클래스) 로드")
    add_bullet(doc, "MediaPipe Hands 싱글톤 (프로세스 당 1회 초기화)")
    add_bullet(doc, "idx_to_label 의 뒤 숫자 제거(예: '다시3' → '다시') 로 베이스 라벨 매핑 (369개 단어)")
    add_bullet(doc, "aggregate_base: 같은 베이스 라벨 확률 합산")
    add_bullet(doc, "FrameMLP 클래스: nn.Module 래퍼 (self.net = nn.Sequential) — state_dict 키 정확히 매칭")

    add_heading(doc, "8.2 DominantTracker (세션 단위)", 2)
    add_para(
        doc,
        "EMA decay 0.97 로 좌/우 손 손목 이동량을 누적해 주로 쓰는 손 판별. "
        "연결별로 독립적이라 WordSession 에 포함.",
    )

    add_heading(doc, "8.3 MotionSeg (모션 세그먼테이션)", 2)
    add_para(
        doc,
        "손 특징 벡터의 이전 프레임 대비 변화량이 임계값을 넘으면 'active' 상태로 "
        "진입, 일정 시간 정지하면 'ready' 로 세그먼트 종료.",
    )
    add_bullet(doc, "thr (모션 임계값): 0.003 — JPEG 손실/10fps 환경 보정")
    add_bullet(doc, "min_f: 10 프레임 (10fps 기준 1초)")
    add_bullet(doc, "max_f: 100 프레임 (10초 타임아웃)")
    add_bullet(doc, "cool: 20 프레임 (2초 정지 → 세그먼트 완료)")
    add_bullet(doc, "smooth_w: 5 (최근 5프레임 모션 평균)")

    add_heading(doc, "8.4 WordSession (연결 당 상태)", 2)
    add_para(
        doc,
        "WebSocket 연결 하나에 WordSession 하나. DominantTracker, MotionSeg, "
        "프레임별 확률 버퍼(seg_probs), target 프레임별 최고 점수(seg_max_target_score) 보유.",
    )

    add_heading(doc, "8.5 /ws/word_recognition WebSocket 엔드포인트", 2)
    add_para(doc, "파일: main/api/learning/learning_routes.py")
    add_para(
        doc,
        "클라이언트 메시지: {type:'frame', image: b64, target: '단어', "
        "category: 'word', subcategory: 'greeting'}. "
        "서버 응답: {type:'prediction', hand_detected, motion, "
        "segment_top3 (세그먼트 완료 시), score, is_passed, target}.",
    )

    add_heading(doc, "8.6 프론트엔드: 시작 버튼 기반 녹화 UX", 2)
    add_para(doc, "파일: frontend/js/word_learn.js, frontend/html/word_learn.html, frontend/css/learning.css")
    add_bullet(doc, "Step 3 진입 시 카메라는 켜지지만 녹화는 시작 안 함")
    add_bullet(doc, "시작 버튼 클릭 → WebSocket 프레임 전송 개시 + 진행바 노출")
    add_bullet(doc, "녹화 중에는 Top-3 영역 비워두고 진행바(10초) + 경과 시간 (X.X / 10.0 초) 표시")
    add_bullet(doc, "서버가 segment_top3 보내오면 진행바 숨기고 Top-3 + 점수 표시")
    add_bullet(doc, "확인 버튼 → DB 저장, 다시 녹화 버튼 → 재시도")
    add_bullet(doc, "10초 초과 시 '시간 초과' 표시 + 재녹화 가능")

    # ================================================================
    # 9. 점수 로직 개선
    # ================================================================
    add_heading(doc, "9. 점수·Top-3 일관성 개선", 1)

    add_heading(doc, "9.1 타겟 변형(variant) 매칭", 2)
    add_para(
        doc,
        "DB 의 '언니/누나' vs 모델의 '언니_누나', '언니(누나)' 등 표기 차이를 흡수. "
        "_target_variants 로 슬래시, 언더스코어, 괄호, 쉼표를 모두 포함하는 변형 리스트 생성 후 "
        "각 변형을 aggregate_base 에서 조회.",
    )

    add_heading(doc, "9.2 target_score: 합산 → 최대값", 2)
    add_para(
        doc,
        "초기 구현은 변형 확률을 모두 더해 min(total, 1.0) 으로 cap — "
        "'실제 Top-3 상위 단어가 17% 인데 점수가 100% 나오는' 불일치 발생. "
        "변형 중 가장 강하게 매칭된 하나만 선택하도록 수정.",
    )

    add_heading(doc, "9.3 세그먼트 중 프레임별 최고 점수 추적", 2)
    add_para(
        doc,
        "사용자가 동작 중간에 잠깐이라도 정답에 근접하면 그 값을 점수에 반영. "
        "WordSession.seg_max_target_score 에 매 프레임 최대값 기록. "
        "세그먼트 종료 시 이 값을 최종 점수로 사용.",
    )

    add_heading(doc, "9.4 Top-3 도 '프레임별 최고값' 기준", 2)
    add_para(
        doc,
        "점수는 peak(max), Top-3 는 avg 였던 불일치 해결. "
        "세그먼트 내 각 base label 별 프레임 최고 확률을 기록 → "
        "그 값으로 정렬한 Top-3 반환. 점수와 1위 숫자가 동일하도록 일치시킴.",
    )

    add_heading(doc, "9.5 Top-3 를 타겟 카테고리 내부로 필터링", 2)
    add_para(
        doc,
        "타겟이 'greeting' 카테고리면 Top-3 도 greeting 카테고리 단어로만 제한. "
        "WordRecognitionService.category_allowed 에 각 카테고리 별 허용 라벨 세트(변형 포함) 구축. "
        "fingerspell 확장 9개도 별도 세트로.",
    )

    # ================================================================
    # 10. 인식률 튜닝
    # ================================================================
    add_heading(doc, "10. 인식률/반응성 튜닝", 1)

    add_heading(doc, "10.1 프레임 전송 주기", 2)
    add_para(doc, "FRAME_INTERVAL_MS: 200ms (5fps) → 100ms (10fps) 로 2배 밀도.")

    add_heading(doc, "10.2 JPEG 품질", 2)
    add_para(
        doc,
        "canvas.toDataURL 품질 0.6 → 0.85 — MediaPipe 손 검출 안정성 향상. "
        "keypoint 정확도 저하로 인한 false-idle 감소.",
    )

    add_heading(doc, "10.3 모션 임계값", 2)
    add_para(
        doc,
        "recon_word.py 원본 0.008 → 0.003 으로 추가 인하. "
        "JPEG 압축 손실로 MediaPipe 가 간헐적으로 손을 놓치면 motion 값이 0 으로 깎이는데, "
        "이를 보정.",
    )

    add_heading(doc, "10.4 타이밍 조정 (10fps 기준)", 2)
    add_bullet(doc, "최소 녹화: 1초 (min_f=10)")
    add_bullet(doc, "정지 감지: 2초 (cool=20)")
    add_bullet(doc, "최대 녹화: 10초 (max_f=100 / RECORD_MAX_MS=10000)")

    # ================================================================
    # 11. 남은 과제
    # ================================================================
    add_heading(doc, "11. 남은 과제 / 주의 사항", 1)
    add_bullet(
        doc,
        "웹 환경과 desktop(recon_word.py) 의 프레임 레이트(10 vs 30 fps), "
        "JPEG 압축, 해상도 차이로 인식률은 완전 동일하기 어려움.",
    )
    add_bullet(
        doc,
        "MediaPipe Hands 는 연결 간 공유 인스턴스 — 다중 사용자 동시 접속 시 tracking state 간섭 가능. "
        "_infer_lock 으로 직렬화됨.",
    )
    add_bullet(
        doc,
        "로그인 콜백은 JSON 반환 방식 유지 — 프론트로 자동 리다이렉트는 하지 않음. "
        "사용자가 수동으로 페이지 이동 또는 후속 작업 필요.",
    )
    add_bullet(
        doc,
        "단어 모델 model_word.pt 의 라벨이 훈련 데이터 기준이라 "
        "DB 단어와 표기가 다를 수 있음 (variant 매칭으로 대부분 해결).",
    )

    # 저장
    output = "/Users/garyeong/Desktop/Real-time-sign-language-translation-service-Backend-master/changelog_26_04_19.docx"
    doc.save(output)
    print(f"저장 완료: {output}")


if __name__ == "__main__":
    main()
